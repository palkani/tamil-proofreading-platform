#!/usr/bin/env python3
"""
Web-based Image to Text Extractor Tool
Upload images/PDFs and download extracted text as Word document
"""

from flask import Flask, render_template, request, send_file, jsonify
from werkzeug.utils import secure_filename
import os
from PIL import Image
from pdf2image import convert_from_path
from ocr_engine import run_ocr
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import shutil
from pathlib import Path

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = tempfile.mkdtemp()

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'pdf', 'tiff', 'bmp', 'gif'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_image(image_path, lang='eng+tam'):
    """Extract text from image using OCR (native C++ module or pytesseract) with Tamil support"""
    try:
        image = Image.open(image_path)
        text = run_ocr(image, lang=lang)
        return text
    except Exception as e:
        if 'tam' in lang:
            try:
                text = run_ocr(image, lang='eng')
                return text
            except Exception:
                pass
        return f"Error: {str(e)}"

def extract_text_from_pdf(pdf_path, lang='eng+tam'):
    """Extract text from PDF with Tamil support"""
    try:
        images = convert_from_path(pdf_path, dpi=300)
        all_text = []
        
        for i, image in enumerate(images, 1):
            try:
                text = run_ocr(image, lang=lang)
            except Exception:
                text = run_ocr(image, lang='eng')
            all_text.append(f"--- Page {i} ---\n{text}\n")
        
        return "\n".join(all_text)
    except Exception as e:
        return f"Error: {str(e)}"

def create_word_document(text, output_path, original_filename):
    """Create Word document with extracted text"""
    doc = Document()
    
    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run(f"Text Extracted from: {original_filename}")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("_" * 80)
    
    # Add content
    doc.add_paragraph(text)
    
    # Save
    doc.save(output_path)

@app.route('/')
def index():
    """Main page"""
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload and OCR processing"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'File type not allowed'}), 400
    
    try:
        # Save uploaded file
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Extract text
        file_ext = Path(filename).suffix.lower()
        
        # Get language preference from request (default to Tamil + English)
        lang = request.form.get('lang', 'eng+tam')
        
        if file_ext == '.pdf':
            extracted_text = extract_text_from_pdf(filepath, lang=lang)
        else:
            extracted_text = extract_text_from_image(filepath, lang=lang)
        
        # Create Word document
        output_filename = f"{Path(filename).stem}_extracted.docx"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
        
        create_word_document(extracted_text, output_path, filename)
        
        # Clean up input file
        os.remove(filepath)
        
        return jsonify({
            'success': True,
            'text': extracted_text[:500] + ('...' if len(extracted_text) > 500 else ''),
            'full_text': extracted_text,  # Include full text for frontend
            'download_filename': output_filename,
            'char_count': len(extracted_text)
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download/<filename>')
def download_file(filename):
    """Download the generated Word document"""
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    return send_file(filepath, as_attachment=True, download_name=filename)

if __name__ == '__main__':
    print("\n" + "="*60)
    print("🚀 OCR Tool - Image to Text Converter")
    print("="*60)
    print("\n📱 Open your browser and go to: http://localhost:5000")
    print("\n⌨️  Press Ctrl+C to stop the server\n")
    
    app.run(debug=True, host='0.0.0.0', port=5000)
