#!/usr/bin/env python3
"""
Image to Text Extractor Tool
Extracts text from images (JPEG, PNG) and PDFs, then converts to Word format
"""

import os
import sys
from PIL import Image
from pdf2image import convert_from_path
from ocr_engine import run_ocr
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import argparse
from pathlib import Path

class ImageToTextConverter:
    """
    Convert images and PDFs to text and save as Word documents
    """
    
    def __init__(self):
        self.supported_image_formats = ['.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif']
        self.supported_formats = self.supported_image_formats + ['.pdf']
    
    def extract_text_from_image(self, image_path):
        """
        Extract text from a single image using OCR
        
        Args:
            image_path: Path to the image file
            
        Returns:
            Extracted text as string
        """
        try:
            # Open image
            image = Image.open(image_path)
            # Perform OCR (native C++ or pytesseract)
            text = run_ocr(image)
            return text
        except Exception as e:
            print(f"Error extracting text from image: {e}")
            return ""
    
    def extract_text_from_pdf(self, pdf_path):
        """
        Extract text from PDF by converting pages to images
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Extracted text as string
        """
        try:
            print(f"Converting PDF to images...")
            
            # Convert PDF pages to images
            images = convert_from_path(pdf_path, dpi=300)
            
            all_text = []
            
            # Process each page
            for i, image in enumerate(images, 1):
                print(f"Processing page {i} of {len(images)}...")
                text = run_ocr(image)
                all_text.append(f"--- Page {i} ---\n{text}\n")
            
            return "\n".join(all_text)
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""
    
    def create_word_document(self, text, output_path, title="Extracted Text"):
        """
        Create a Word document with the extracted text
        
        Args:
            text: Extracted text to put in document
            output_path: Path where to save the Word document
            title: Title for the document
        """
        try:
            # Create document
            doc = Document()
            
            # Add title
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(18)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add separator
            doc.add_paragraph("_" * 80)
            
            # Add extracted text
            doc.add_paragraph(text)
            
            # Save document
            doc.save(output_path)
            print(f"\n✅ Word document created: {output_path}")
            
        except Exception as e:
            print(f"Error creating Word document: {e}")
    
    def convert_file(self, input_path, output_path=None):
        """
        Main conversion function
        
        Args:
            input_path: Path to input file (image or PDF)
            output_path: Path for output Word file (optional)
        """
        # Validate input file
        input_file = Path(input_path)
        
        if not input_file.exists():
            print(f"❌ Error: File not found: {input_path}")
            return
        
        file_ext = input_file.suffix.lower()
        
        if file_ext not in self.supported_formats:
            print(f"❌ Error: Unsupported file format: {file_ext}")
            print(f"Supported formats: {', '.join(self.supported_formats)}")
            return
        
        # Set output path if not provided
        if output_path is None:
            output_path = input_file.stem + "_extracted.docx"
        
        print(f"\n🔍 Processing: {input_path}")
        print(f"📄 Output will be: {output_path}\n")
        
        # Extract text based on file type
        if file_ext == '.pdf':
            extracted_text = self.extract_text_from_pdf(input_path)
        else:
            extracted_text = self.extract_text_from_image(input_path)
        
        if extracted_text.strip():
            # Create Word document
            self.create_word_document(
                extracted_text, 
                output_path, 
                title=f"Text Extracted from: {input_file.name}"
            )
            
            print(f"\n✨ Success! Text extracted and saved.")
            print(f"📊 Total characters extracted: {len(extracted_text)}")
        else:
            print("⚠️  Warning: No text was extracted from the file.")
            print("This could mean:")
            print("  - The image quality is too low")
            print("  - The image doesn't contain text")
            print("  - The text is in an unsupported language")


def main():
    """
    Command-line interface for the OCR tool
    """
    parser = argparse.ArgumentParser(
        description='Extract text from images and PDFs, convert to Word format',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python ocr_tool.py image.jpg
  python ocr_tool.py document.pdf -o extracted.docx
  python ocr_tool.py screenshot.png -o output.docx
        """
    )
    
    parser.add_argument('input', help='Input file (image or PDF)')
    parser.add_argument('-o', '--output', help='Output Word file path (optional)')
    
    args = parser.parse_args()
    
    # Create converter and process file
    converter = ImageToTextConverter()
    converter.convert_file(args.input, args.output)


if __name__ == "__main__":
    main()
