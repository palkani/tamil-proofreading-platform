#!/usr/bin/env python3
"""
Document Converter API - ProofTamil
Converts between multiple document formats
Supports: DOCX, PDF, TXT, RTF, HTML, ODT
"""

from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import tempfile
import shutil
from pathlib import Path
from datetime import datetime
import subprocess
import docx2pdf
from pypdf import PdfReader, PdfWriter
from docx import Document
from docx.shared import Pt, Inches
import mammoth
import pypandoc

app = Flask(__name__)
CORS(app)  # Enable CORS for frontend integration

app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max
app.config['UPLOAD_FOLDER'] = tempfile.mkdtemp()

# Supported conversions
CONVERSIONS = {
    'pdf': ['docx', 'txt', 'html', 'rtf'],
    'docx': ['pdf', 'txt', 'html', 'rtf', 'odt'],
    'txt': ['pdf', 'docx', 'html', 'rtf'],
    'html': ['pdf', 'docx', 'txt', 'rtf'],
    'rtf': ['pdf', 'docx', 'txt', 'html'],
    'odt': ['pdf', 'docx', 'txt', 'html']
}

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'html', 'rtf', 'odt', 'doc'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_file_info(filepath):
    """Get file information"""
    stat = os.stat(filepath)
    return {
        'size': stat.st_size,
        'size_mb': round(stat.st_size / (1024 * 1024), 2),
        'modified': datetime.fromtimestamp(stat.st_mtime).isoformat()
    }

# ============= CONVERSION FUNCTIONS =============

def pdf_to_docx(input_path, output_path):
    """Convert PDF to DOCX"""
    try:
        # Extract text from PDF
        reader = PdfReader(input_path)
        
        # Create new Word document
        doc = Document()
        
        # Add title
        title = doc.add_paragraph()
        title_run = title.add_run("Converted from PDF")
        title_run.bold = True
        title_run.font.size = Pt(16)
        
        doc.add_paragraph("_" * 80)
        
        # Extract and add text from each page
        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                doc.add_heading(f"Page {i}", level=2)
                doc.add_paragraph(text)
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"PDF to DOCX error: {e}")
        return False

def pdf_to_txt(input_path, output_path):
    """Convert PDF to TXT"""
    try:
        reader = PdfReader(input_path)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            for i, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                f.write(f"=== Page {i} ===\n\n")
                f.write(text)
                f.write("\n\n")
        
        return True
    except Exception as e:
        print(f"PDF to TXT error: {e}")
        return False

def docx_to_pdf(input_path, output_path):
    """Convert DOCX to PDF"""
    try:
        # Use LibreOffice for conversion (best quality)
        subprocess.run([
            'soffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', os.path.dirname(output_path),
            input_path
        ], check=True, capture_output=True)
        
        # LibreOffice creates file with same name but .pdf extension
        temp_pdf = os.path.join(
            os.path.dirname(output_path),
            Path(input_path).stem + '.pdf'
        )
        
        if os.path.exists(temp_pdf) and temp_pdf != output_path:
            shutil.move(temp_pdf, output_path)
        
        return os.path.exists(output_path)
    except Exception as e:
        print(f"DOCX to PDF error: {e}")
        return False

def docx_to_txt(input_path, output_path):
    """Convert DOCX to TXT"""
    try:
        doc = Document(input_path)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                f.write(para.text + '\n')
        
        return True
    except Exception as e:
        print(f"DOCX to TXT error: {e}")
        return False

def docx_to_html(input_path, output_path):
    """Convert DOCX to HTML"""
    try:
        with open(input_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value
        
        # Wrap in proper HTML structure
        full_html = f"""
<!DOCTYPE html>
<html lang="ta">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Converted Document</title>
    <style>
        body {{
            font-family: 'Latha', 'Tamil MN', sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
        }}
    </style>
</head>
<body>
    {html}
</body>
</html>
"""
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_html)
        
        return True
    except Exception as e:
        print(f"DOCX to HTML error: {e}")
        return False

def txt_to_docx(input_path, output_path):
    """Convert TXT to DOCX"""
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        doc = Document()
        
        # Add title
        title = doc.add_paragraph()
        title_run = title.add_run("Converted from Text")
        title_run.bold = True
        title_run.font.size = Pt(16)
        
        doc.add_paragraph("_" * 80)
        
        # Add text with paragraphs
        for para in text.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"TXT to DOCX error: {e}")
        return False

def txt_to_pdf(input_path, output_path):
    """Convert TXT to PDF via DOCX"""
    try:
        # First convert to DOCX
        temp_docx = output_path.replace('.pdf', '.docx')
        if txt_to_docx(input_path, temp_docx):
            # Then convert DOCX to PDF
            result = docx_to_pdf(temp_docx, output_path)
            os.remove(temp_docx)
            return result
        return False
    except Exception as e:
        print(f"TXT to PDF error: {e}")
        return False

def html_to_docx(input_path, output_path):
    """Convert HTML to DOCX using pandoc"""
    try:
        pypandoc.convert_file(
            input_path,
            'docx',
            outputfile=output_path,
            extra_args=['--reference-doc=/dev/null']
        )
        return True
    except Exception as e:
        print(f"HTML to DOCX error: {e}")
        return False

def generic_conversion(input_path, output_path, from_format, to_format):
    """Generic conversion using pandoc"""
    try:
        pypandoc.convert_file(
            input_path,
            to_format,
            outputfile=output_path,
            format=from_format
        )
        return True
    except Exception as e:
        print(f"Generic conversion error ({from_format} to {to_format}): {e}")
        return False

def convert_document(input_path, output_path, from_format, to_format):
    """Main conversion dispatcher"""
    conversion_key = f"{from_format}_to_{to_format}"
    
    # Map of specific conversion functions
    converters = {
        'pdf_to_docx': pdf_to_docx,
        'pdf_to_txt': pdf_to_txt,
        'docx_to_pdf': docx_to_pdf,
        'docx_to_txt': docx_to_txt,
        'docx_to_html': docx_to_html,
        'txt_to_docx': txt_to_docx,
        'txt_to_pdf': txt_to_pdf,
        'html_to_docx': html_to_docx,
    }
    
    # Use specific converter if available
    if conversion_key in converters:
        return converters[conversion_key](input_path, output_path)
    
    # Otherwise use generic pandoc conversion
    return generic_conversion(input_path, output_path, from_format, to_format)

# ============= API ENDPOINTS =============

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'service': 'Document Converter API',
        'version': '1.0.0'
    })

@app.route('/api/supported-conversions', methods=['GET'])
def get_supported_conversions():
    """Get list of supported conversions"""
    return jsonify({
        'conversions': CONVERSIONS,
        'formats': list(CONVERSIONS.keys())
    })

@app.route('/api/convert', methods=['POST'])
def convert_file():
    """Convert document endpoint"""
    
    # Check if file is present
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    to_format = request.form.get('to_format', '').lower()
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not to_format:
        return jsonify({'error': 'Target format not specified'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'File type not supported'}), 400
    
    try:
        # Get file info
        filename = secure_filename(file.filename)
        from_format = Path(filename).suffix[1:].lower()
        
        # Validate conversion is supported
        if from_format not in CONVERSIONS:
            return jsonify({'error': f'Source format {from_format} not supported'}), 400
        
        if to_format not in CONVERSIONS.get(from_format, []):
            return jsonify({'error': f'Cannot convert {from_format} to {to_format}'}), 400
        
        # Save uploaded file
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)
        
        # Get input file info
        input_info = get_file_info(input_path)
        
        # Prepare output path
        output_filename = f"{Path(filename).stem}_converted.{to_format}"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
        
        # Perform conversion
        success = convert_document(input_path, output_path, from_format, to_format)
        
        if not success or not os.path.exists(output_path):
            return jsonify({'error': 'Conversion failed'}), 500
        
        # Get output file info
        output_info = get_file_info(output_path)
        
        # Clean up input file
        os.remove(input_path)
        
        return jsonify({
            'success': True,
            'message': f'Successfully converted {from_format.upper()} to {to_format.upper()}',
            'download_filename': output_filename,
            'input_size_mb': input_info['size_mb'],
            'output_size_mb': output_info['size_mb'],
            'from_format': from_format.upper(),
            'to_format': to_format.upper()
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/download/<filename>', methods=['GET'])
def download_file(filename):
    """Download converted file"""
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    
    if not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404
    
    return send_file(
        filepath,
        as_attachment=True,
        download_name=filename
    )

@app.route('/api/cleanup', methods=['POST'])
def cleanup():
    """Cleanup temporary files"""
    try:
        # Remove old files (older than 1 hour)
        import time
        current_time = time.time()
        
        for filename in os.listdir(app.config['UPLOAD_FOLDER']):
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            if os.path.isfile(filepath):
                file_age = current_time - os.path.getmtime(filepath)
                if file_age > 3600:  # 1 hour
                    os.remove(filepath)
        
        return jsonify({'success': True, 'message': 'Cleanup completed'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    print("\n" + "="*70)
    print("🔄 Document Converter API - ProofTamil")
    print("="*70)
    print("\n📡 API running on: http://localhost:5001")
    print("\n📚 Endpoints:")
    print("   POST /api/convert - Convert documents")
    print("   GET  /api/download/<filename> - Download converted file")
    print("   GET  /api/supported-conversions - Get supported formats")
    print("\n⌨️  Press Ctrl+C to stop\n")
    
    app.run(debug=True, host='0.0.0.0', port=5001)
